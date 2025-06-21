from typing import List, BinaryIO
from docx import Document
from io import BytesIO
from .base import BaseParser

class DocxParser(BaseParser):
    """Parser for Microsoft Word documents."""
    
    def parse(self, file: BinaryIO) -> List[str]:
        """
        Parse a Word document and return a list of text chunks.
        
        Args:
            file (BinaryIO): Word document file object
            
        Returns:
            List[str]: List of text chunks
        """
        # Create document object
        doc = Document(BytesIO(file.read()))
        
        # Extract text from paragraphs and tables
        text_content = []
        
        # Get text from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_content.append(text)
        
        # Get text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        # Join all content and chunk the text
        full_text = '\n'.join(text_content)
        return self.chunk_text(self.clean_text(full_text)) 